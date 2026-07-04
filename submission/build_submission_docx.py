from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission"


BANK_CONTEXT = (
    "Context: an anonymized bank is treated as an early-stage bank. The design favours fast, "
    "boring delivery paths with reliability controls that protect customer trust, "
    "account-read availability, and operational speed."
)


def read(path: str, max_lines: int | None = None) -> str:
    text = (ROOT / path).read_text()
    if max_lines is None:
        return text.rstrip()
    return "\n".join(text.splitlines()[:max_lines]).rstrip()


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(18, 47, 87)

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor(35, 35, 35)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.space_after = Pt(0)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Bank Account Service Platform")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(18, 47, 87)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)

    doc.add_paragraph(subtitle)
    doc.add_paragraph(BANK_CONTEXT)
    doc.add_paragraph(
        "Repository path: /home/imich/learning/FinTechProdService"
    )


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_evidence_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Evidence"
    hdr[1].text = "Why it matters"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def add_code(doc: Document, path: str, max_lines: int | None = None) -> None:
    doc.add_heading(path, level=3)
    for line in read(path, max_lines).splitlines():
        doc.add_paragraph(line if line else " ", style="CodeBlock")


def save_doc(name: str, title: str, subtitle: str, sections) -> None:
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_cover(doc, title, subtitle)
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        body(doc)

    doc.save(OUT / name)


def postgres_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "Postgres is provisioned as AWS RDS in Terraform for production shape. "
            "Local execution uses Docker Compose Postgres only to keep kind runnable "
            "without cloud resources."
        )
        add_evidence_table(
            doc,
            [
                ("infra/terraform/modules/rds/main.tf", "RDS Postgres, encrypted storage, backups, deletion protection, IAM DB auth."),
                ("local/postgres/002_runtime_role.sql", "Least-privilege runtime role and explicit grants."),
                ("local/postgres/004_runtime_role_password.sh", "Local runtime password injected from environment, not committed."),
                ("docs/backup-restore.md", "RDS PITR, RPO/RTO, and restore validation thinking."),
            ],
        )

    def reliability(doc: Document) -> None:
        add_bullets(
            doc,
            [
                "Fast for an early bank: Docker Compose keeps local development simple.",
                "Reliable for production shape: RDS gives managed backups, PITR, encryption, and operational maturity.",
                "Blast radius is constrained because the application role is not a database owner and has no DDL permissions.",
                "Restore is treated as a tested workflow, not just a checkbox.",
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "local/postgres/002_runtime_role.sql")
        add_code(doc, "local/postgres/004_runtime_role_password.sh")
        add_code(doc, "infra/terraform/modules/rds/main.tf", 90)
        add_code(doc, "docs/backup-restore.md")

    save_doc(
        "01_postgres_provisioning_least_privilege.docx",
        "Postgres Provisioning and Least-Privilege Role",
        "SQL / Terraform evidence for RDS production mapping and local least-privilege grants.",
        [
            ("Summary", summary),
            ("Banking Reliability Rationale", reliability),
            ("Evidence Excerpts", excerpts),
        ],
    )


def credentials_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "No plaintext database credential is committed. Local secrets are supplied "
            "through environment variables and Kubernetes Secrets created at runtime. "
            "Production maps to AWS Secrets Manager, External Secrets, IRSA/EKS Pod "
            "Identity, and RDS IAM database authentication."
        )
        add_evidence_table(
            doc,
            [
                (".env.example", "Variable names only, no values."),
                ("k8s/helm/bank-account-service/templates/externalsecret.yaml", "Secrets Manager to Kubernetes Secret path."),
                ("k8s/helm/bank-account-service/templates/serviceaccount.yaml", "IRSA annotation for production service account."),
                ("infra/terraform/modules/irsa/main.tf", "Scoped IAM access to secret, queue, and rds-db:connect."),
                ("infra/terraform/modules/rds/main.tf", "RDS IAM auth enabled."),
                ("docs/database-access.md", "Credential handling narrative."),
            ],
        )

    def flow(doc: Document) -> None:
        add_numbered(
            doc,
            [
                "Developer sets local secrets in shell or ignored .env file.",
                "Docker Compose uses those variables to bootstrap local Postgres.",
                "kind receives DATABASE_URL through a Kubernetes Secret created at runtime.",
                "Production stores database secret material in AWS Secrets Manager.",
                "External Secrets Operator writes the Kubernetes Secret.",
                "IRSA or EKS Pod Identity gives the pod scoped secret read access without static keys.",
                "RDS IAM database authentication is enabled for the production database shape.",
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, ".env.example")
        add_code(doc, "local/docker-compose.yml", 70)
        add_code(doc, "k8s/helm/bank-account-service/templates/externalsecret.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/serviceaccount.yaml")
        add_code(doc, "infra/terraform/modules/irsa/main.tf")
        add_code(doc, "docs/database-access.md")

    save_doc(
        "02_credential_handling_no_plaintext.docx",
        "Credential Handling",
        "No plaintext secrets in repo; local runtime secrets and AWS production credential mapping.",
        [
            ("Summary", summary),
            ("Credential Flow", flow),
            ("Evidence Excerpts", excerpts),
        ],
    )


def messaging_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "The consumer uses a processed_events table with event_id as the primary key. "
            "The side effect is written only after the idempotency insert succeeds inside "
            "the same database transaction."
        )
        add_evidence_table(
            doc,
            [
                ("app/src/bank_account_service/consumer.py", "Transaction, validation, audit side effect."),
                ("app/src/bank_account_service/idempotency.py", "ON CONFLICT DO NOTHING idempotency guard."),
                ("app/tests/test_consumer_idempotency.py", "Duplicate event creates one audit row."),
                ("infra/terraform/modules/sqs/main.tf", "SQS main queue, DLQ, maxReceiveCount = 3."),
                ("docs/messaging-semantics.md", "Poison message and retry semantics."),
            ],
        )

    def business(doc: Document) -> None:
        add_bullets(
            doc,
            [
                "Duplicate messages must not duplicate customer-visible audit side effects.",
                "Transient database failures are retried by SQS visibility timeout behaviour.",
                "Malformed messages move to DLQ after maxReceiveCount = 3 so engineers can inspect and replay deliberately.",
                "SQS is enough for this early bank because this side effect does not need global ordering or stream replay.",
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "app/src/bank_account_service/consumer.py")
        add_code(doc, "app/src/bank_account_service/idempotency.py")
        add_code(doc, "app/tests/test_consumer_idempotency.py")
        add_code(doc, "infra/terraform/modules/sqs/main.tf")
        add_code(doc, "docs/messaging-semantics.md")

    save_doc(
        "03_idempotent_consumer_dlq.docx",
        "Idempotent Message Consumer and DLQ Semantics",
        "Consumer implementation with poison message handling and dead-letter routing.",
        [
            ("Summary", summary),
            ("Business Reliability Rationale", business),
            ("Evidence Excerpts", excerpts),
        ],
    )


def cicd_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "CI/CD is split by responsibility under .github/workflows: Docker, "
            "Kubernetes, Terraform plan, and Terraform apply. Reusable actions are "
            "pinned by commit SHA, images are tagged by Git SHA, and production-style "
            "deploy/apply jobs require a GitHub Environment manual gate."
        )
        add_evidence_table(
            doc,
            [
                (".github/workflows/docker.yml", "Tests, Docker build, Trivy HIGH/CRITICAL gate, immutable image tag."),
                (".github/workflows/kubernetes.yml", "Helm lint/template and manual production deploy with OIDC."),
                (".github/workflows/terraform-plan.yml", "Terraform fmt/init/validate and optional OIDC plan."),
                (".github/workflows/terraform-apply.yml", "Manual production-gated apply."),
                ("docs/rollback.md", "Helm revision and immutable SHA rollback paths."),
            ],
        )

    def controls(doc: Document) -> None:
        add_bullets(
            doc,
            [
                "Image tags use GITHUB_SHA, not latest.",
                "Trivy fails the Docker workflow on HIGH or CRITICAL image vulnerabilities.",
                "OIDC is used for cloud authentication; static cloud keys are not committed.",
                "Production deploy/apply jobs use GitHub Environments for approval.",
                "Rollback can target either a Helm revision or a previous good immutable SHA.",
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, ".github/workflows/docker.yml")
        add_code(doc, ".github/workflows/kubernetes.yml")
        add_code(doc, ".github/workflows/terraform-plan.yml")
        add_code(doc, ".github/workflows/terraform-apply.yml")
        add_code(doc, "docs/rollback.md")

    save_doc(
        "04_github_actions_cicd.docx",
        "GitHub Actions CI/CD",
        "Split pipelines with scan gate, immutable tag, OIDC cloud auth, manual approval, and rollback.",
        [
            ("Summary", summary),
            ("Controls", controls),
            ("Evidence Excerpts", excerpts),
        ],
    )


def terraform_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "Terraform expresses the AWS production mapping only. It is formatted and "
            "validated locally with backend disabled, but not applied for the assessment."
        )
        add_evidence_table(
            doc,
            [
                ("infra/terraform/main.tf", "Root module calls ecr, eks, irsa, rds, secrets, sqs, observability."),
                ("infra/terraform/modules/ecr", "ECR immutable tags and scan on push."),
                ("infra/terraform/modules/eks", "EKS cluster and managed node group shape."),
                ("infra/terraform/modules/rds", "Encrypted RDS Postgres, backups, deletion protection, IAM DB auth."),
                ("infra/terraform/modules/secrets", "Secret metadata, no hardcoded secret value."),
                ("infra/terraform/plans/sanitized-plan.txt", "Representative plan summary with no account IDs or secrets."),
            ],
        )

    def tradeoff(doc: Document) -> None:
        add_bullets(
            doc,
            [
                "Simple modules make the design reviewable in a short assessment.",
                "No live AWS data sources are used, so validate does not require cloud discovery.",
                "Placeholder VPC/subnet/account values show intent without pretending this was deployed.",
                "This is enough to discuss production shape while keeping local kind independent.",
            ],
        )

    def relationships(doc: Document) -> None:
        doc.add_paragraph(
            "The AWS resources are not isolated boxes; they support each other in a "
            "small production platform."
        )
        add_evidence_table(
            doc,
            [
                ("ECR and EKS", "ECR stores the approved image. EKS pulls that immutable image and runs it as pods."),
                ("EKS and RDS", "The service runs on EKS and reads durable account data from RDS Postgres."),
                ("Secrets Manager, External Secrets, and IRSA", "Secrets stay outside Git; the pod gets only the secret access it needs."),
                ("SQS and DLQ", "SQS keeps side effects off the customer request path; DLQ quarantines repeated bad messages."),
                ("Observability", "Metrics, dashboards, and alerts show whether account reads are reliable for customers."),
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "infra/terraform/main.tf")
        add_code(doc, "infra/terraform/modules/ecr/main.tf")
        add_code(doc, "infra/terraform/modules/eks/main.tf")
        add_code(doc, "infra/terraform/modules/rds/main.tf")
        add_code(doc, "infra/terraform/modules/irsa/main.tf")
        add_code(doc, "infra/terraform/modules/secrets/main.tf")
        add_code(doc, "infra/terraform/plans/sanitized-plan.txt")

    save_doc(
        "05_terraform_cloud_resources_plan.docx",
        "Terraform Cloud Resources and Sanitized Plan",
        "AWS production mapping for EKS, IAM, ECR, RDS, secrets, SQS, and observability.",
        [
            ("Summary", summary),
            ("Trade-offs", tradeoff),
            ("Resource Relationships For Non-Technical Reviewers", relationships),
            ("Evidence Excerpts", excerpts),
        ],
    )


def observability_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "Observability uses RED metrics, one trace path through account lookup, "
            "a simple dashboard, a business-impacting SLO burn alert, and PII scrubbing notes."
        )
        add_evidence_table(
            doc,
            [
                ("app/src/bank_account_service/metrics.py", "Request count, errors, and duration metrics."),
                ("observability/traces/trace-path.md", "Account lookup trace path without raw account ID telemetry."),
                ("observability/dashboards/bank-account-service-dashboard.json", "Request rate, error rate, p95 latency panels."),
                ("observability/alerts/bank-account-service-alerts.yaml", "SLO burn alert with business impact annotation."),
                ("observability/data-scrubbing.md", "PII and high-cardinality telemetry controls."),
            ],
        )

    def business(doc: Document) -> None:
        add_bullets(
            doc,
            [
                "Request rate tells the bank whether demand changed during onboarding, campaigns, releases, or incidents.",
                "Error rate is customer-impacting: account reads failing create trust and support risk.",
                "P95 latency protects customer experience where averages hide slow banking journeys.",
                "The page alert is tied to account lookup SLO burn, not CPU alone.",
                "PII scrubbing prevents observability from becoming a data exposure path.",
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "app/src/bank_account_service/metrics.py", 120)
        add_code(doc, "observability/alerts/bank-account-service-alerts.yaml")
        add_code(doc, "observability/traces/trace-path.md")
        add_code(doc, "observability/dashboards/bank-account-service-dashboard.json", 120)
        add_code(doc, "observability/data-scrubbing.md")

    save_doc(
        "06_observability_business_o11y.docx",
        "Observability Configuration",
        "RED metrics, trace path, dashboard panels, alert rule, and PII scrubbing notes.",
        [
            ("Summary", summary),
            ("Business Context", business),
            ("Evidence Excerpts", excerpts),
        ],
    )


def readme_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "README.md is the reviewer entry point. It shows how to bring up local "
            "dependencies, build/load/deploy to kind, validate Helm and Terraform, "
            "and map local components to AWS production services."
        )
        add_evidence_table(
            doc,
            [
                ("README.md", "Run path, validation commands, requirement map, local-to-AWS mapping."),
                ("Makefile", "Short commands matching README."),
                ("k8s/helm/bank-account-service/values.yaml", "kind-friendly defaults without committed secrets."),
                ("k8s/helm/bank-account-service/values-prod.yaml", "AWS production image/secret/controller mapping."),
            ],
        )

    def path(doc: Document) -> None:
        add_numbered(
            doc,
            [
                "Set local-only secret values in shell or ignored .env.",
                "Run make local-up for Docker Compose dependencies.",
                "Build image with app/Dockerfile.",
                "Create kind cluster and load local image.",
                "Create Kubernetes database Secret at runtime.",
                "Deploy Helm chart.",
                "Curl health, readiness, metrics, and account lookup endpoints.",
            ],
        )

    def relationships(doc: Document) -> None:
        add_evidence_table(
            doc,
            [
                ("Docker image and ECR", "CI builds, scans, and tags the image; ECR stores the approved production artifact."),
                ("ECR and EKS", "EKS pulls the exact approved image from ECR to run the service."),
                ("EKS and readiness", "EKS only routes traffic to pods that are alive, ready, and not draining."),
                ("RDS and account service", "RDS stores the account data and idempotency records the service depends on."),
                ("SQS and customer speed", "SQS lets slower audit side effects happen asynchronously so account reads stay fast."),
                ("DLQ and operations", "The DLQ isolates poison messages so support and engineering can inspect them without stopping normal flow."),
                ("Observability and trust", "RED metrics and SLO alerts tell the team when customer-visible banking reads are failing or slow."),
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "README.md")
        add_code(doc, "Makefile")
        add_code(doc, "k8s/helm/bank-account-service/values.yaml")
        add_code(doc, "k8s/helm/bank-account-service/values-prod.yaml")

    save_doc(
        "07_readme_platform_mapping.docx",
        "README and Platform Mapping",
        "How to bring the platform up locally and how local services map to AWS.",
        [
            ("Summary", summary),
            ("Review Path", path),
            ("Architecture In Plain English", relationships),
            ("Evidence Excerpts", excerpts),
        ],
    )


def kubernetes_doc() -> None:
    def summary(doc: Document) -> None:
        doc.add_paragraph(
            "The Helm chart provides the Kubernetes workload controls requested for "
            "the assessment: Deployment, Service, probes, rolling-update strategy, "
            "PodDisruptionBudget, security contexts, NetworkPolicy, ServiceAccount, "
            "and HPA. The defaults run on kind without AWS controllers; production "
            "values turn on AWS-style identity and external secrets."
        )
        add_evidence_table(
            doc,
            [
                ("k8s/helm/bank-account-service/templates/deployment.yaml", "Deployment, probes, zero-downtime rollout, graceful drain, resources, security context."),
                ("k8s/helm/bank-account-service/templates/service.yaml", "Stable in-cluster address that routes only to ready pods."),
                ("k8s/helm/bank-account-service/templates/pdb.yaml", "Keeps at least one pod available during voluntary disruption."),
                ("k8s/helm/bank-account-service/templates/hpa.yaml", "Simple CPU autoscaling for the HTTP service."),
                ("k8s/helm/bank-account-service/templates/networkpolicy.yaml", "Ingress/egress intent and blast-radius reduction."),
                ("k8s/helm/bank-account-service/templates/serviceaccount.yaml", "Workload identity hook for production."),
                ("k8s/helm/bank-account-service/values.yaml", "kind-friendly local defaults."),
                ("k8s/helm/bank-account-service/values-prod.yaml", "AWS production-style overrides."),
            ],
        )

    def bank_rationale(doc: Document) -> None:
        add_evidence_table(
            doc,
            [
                ("Deployment rolling update", "For a young bank, ordinary releases should not interrupt account reads. maxUnavailable=0 keeps old pods serving while a new pod becomes ready."),
                ("Readiness probe", "Protects customers by removing pods that are draining or cannot reach Postgres from the Service endpoints."),
                ("Liveness probe", "Keeps dead processes from lingering, but does not depend on Postgres, avoiding restart storms during database incidents."),
                ("preStop + termination grace", "Gives load-balancer and endpoint propagation time so in-flight requests can complete."),
                ("Service", "Provides a stable name for callers and only routes to ready pods."),
                ("PodDisruptionBudget", "Prevents voluntary maintenance from taking every replica down at once."),
                ("Security context", "Cheap, high-value hardening: non-root user, no privilege escalation, read-only root filesystem, and dropped Linux capabilities."),
                ("HPA", "A simple first autoscaler for the HTTP workload. KEDA would be the next step for queue-depth scaling if the consumer became a separate deployment."),
                ("NetworkPolicy", "Documents network intent early and starts shrinking blast radius before the platform grows complex."),
            ],
        )

    def excerpts(doc: Document) -> None:
        add_code(doc, "k8s/helm/bank-account-service/templates/deployment.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/service.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/pdb.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/hpa.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/networkpolicy.yaml")
        add_code(doc, "k8s/helm/bank-account-service/templates/serviceaccount.yaml")
        add_code(doc, "k8s/helm/bank-account-service/values.yaml")
        add_code(doc, "k8s/helm/bank-account-service/values-prod.yaml")

    save_doc(
        "08_kubernetes_helm_controls.docx",
        "Kubernetes Manifests and Helm Chart",
        "Deployment, Service, probes, zero-downtime rollout, PDB, securityContext, NetworkPolicy, and HPA.",
        [
            ("Summary", summary),
            ("Why These Components Help An Early-Stage Bank", bank_rationale),
            ("YAML Evidence", excerpts),
        ],
    )


def main() -> None:
    OUT.mkdir(exist_ok=True)
    postgres_doc()
    credentials_doc()
    messaging_doc()
    cicd_doc()
    terraform_doc()
    observability_doc()
    readme_doc()
    kubernetes_doc()


if __name__ == "__main__":
    main()
